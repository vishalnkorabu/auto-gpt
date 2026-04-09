from __future__ import annotations

from io import BytesIO


def export_report_as_docx(title: str, body: str) -> bytes:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:  # pragma: no cover - depends on local install state.
        raise RuntimeError("DOCX export requires python-docx to be installed.") from exc

    document = Document()
    document.add_heading(title, level=0)
    for block in _paragraphs(body):
        document.add_paragraph(block)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def export_report_as_pdf(title: str, body: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfbase.pdfmetrics import stringWidth
        from reportlab.pdfgen import canvas
    except ModuleNotFoundError as exc:  # pragma: no cover - depends on local install state.
        raise RuntimeError("PDF export requires reportlab to be installed.") from exc

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    width, height = letter
    x = 54
    y = height - 54

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(x, y, title[:80])
    y -= 28
    pdf.setFont("Helvetica", 11)

    for paragraph in _paragraphs(body):
        lines = _wrap_text(paragraph, width - 2 * x, "Helvetica", 11, stringWidth)
        for line in lines or [""]:
            if y < 54:
                pdf.showPage()
                y = height - 54
                pdf.setFont("Helvetica", 11)
            pdf.drawString(x, y, line)
            y -= 15
        y -= 8

    pdf.save()
    return output.getvalue()


def _paragraphs(text: str) -> list[str]:
    return [block.strip() for block in text.splitlines() if block.strip()]


def _wrap_text(text: str, max_width: float, font_name: str, font_size: int, string_width) -> list[str]:
    words = text.split()
    if not words:
        return []

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if string_width(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines
